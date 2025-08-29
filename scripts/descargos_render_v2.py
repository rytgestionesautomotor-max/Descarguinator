# -*- coding: utf-8 -*-
"""
Generador de descargos (.docx) – Versión 2
==========================================
- Compatible con plantillas .docx que usan:
    * Placeholders {{CLAVE}} (soporta notación con puntos: a.b.c)
    * Condicionales con anidamiento real: [[IF ...]] ... [[ELIF ...]] ... [[ELSE]] ... [[/IF]]
      - Expresiones soportadas: VAR | !VAR | VAR == 'valor' | VAR != "valor"
- Limpieza de marcadores residuales y colapso de líneas en blanco múltiples.
- CLI con --plantilla, --caso, --salida, y modo --estricto para fallar si quedan placeholders.
- Rutas robustas usando pathlib (scripts/plantillas/casos/salidas).
- Lee el .docx como texto (párrafos + tablas) y genera salida en párrafos simples (preserva contenido, no el formato).
- Filtros de texto básicos en variables: {{VAR|upper}}, {{VAR|lower}}, {{VAR|title}}.

Requisitos:
    pip install python-docx

Notas sobre esta versión (correcciones y mejoras clave):
- Interpretación correcta de **condicionales anidados** mediante un parser con pila
  que maneja [[IF]], [[ELIF]], [[ELSE]] y [[/IF]].
- Evaluación de condicionales **antes** de sustituir variables (evita tocar ramas descartadas).
- Opción **--estricto** para detectar y fallar si quedan placeholders sin resolver.
- Limpieza más robusta de espacios y líneas en blanco.
"""

from __future__ import annotations
import re
import json
from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Tuple, Optional
import argparse
import sys

from docx import Document


# --------------------------------------------------------------------------------------
# Configuración de rutas
# --------------------------------------------------------------------------------------
BASE = Path(__file__).resolve().parent          # .../Descargos/scripts
ROOT = BASE.parent                              # .../Descargos
DIR_PLANTILLAS = ROOT / "plantillas"
DIR_CASOS = ROOT / "casos"
DIR_SALIDAS = ROOT / "salidas"


# --------------------------------------------------------------------------------------
# Utilidades de IO
# --------------------------------------------------------------------------------------
def ensure_dirs() -> None:
    DIR_SALIDAS.mkdir(exist_ok=True)


def load_json(path: Path) -> Dict[str, Any]:
    if not path.exists():
        raise FileNotFoundError(f"No se encontró el JSON de caso en: {path}")
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def doc_to_text(path: Path) -> str:
    """
    Lee un .docx y devuelve todo el contenido de párrafos y tablas como texto plano.
    (Se preservan saltos de párrafo; las tablas se linealizan por filas).
    """
    if not path.exists():
        raise FileNotFoundError(f"No se encontró la plantilla en: {path}")
    d = Document(str(path))
    parts: List[str] = []

    # Párrafos
    for p in d.paragraphs:
        parts.append(p.text)

    # Tablas (cada fila en una línea)
    for table in d.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))

    return "\n".join(parts)


def text_to_doc(text: str, out_path: Path) -> None:
    """Crea un .docx con párrafos a partir de texto plano."""
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# --------------------------------------------------------------------------------------
# Contexto y utilidades
# --------------------------------------------------------------------------------------
def _get_ctx_value(ctx: Dict[str, Any], key: str) -> Any:
    """Obtiene ctx['a.b.c'] soportando notación con puntos."""
    if "." not in key:
        return ctx.get(key)
    cur: Any = ctx
    for part in key.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return None
    return cur


def _truthy(val: Any) -> bool:
    if isinstance(val, bool):
        return val
    if val is None:
        return False
    if isinstance(val, (int, float)):
        return val != 0
    if isinstance(val, str):
        v = val.strip().lower()
        if v in ("true", "sí", "si", "1", "y", "yes"):
            return True
        if v in ("false", "no", "0", "", "null", "none"):
            return False
        return True
    return True


def _parse_literal(token: str) -> Any:
    """Convierte 'true'/'false'/'null'/'123'/'123.4'/'texto' a python."""
    t = token.strip()
    tl = t.lower()
    if tl in ("true", "false"):
        return tl == "true"
    if tl in ("null", "none"):
        return None
    # Comillas simples o dobles
    if (t.startswith("'") and t.endswith("'")) or (t.startswith('"') and t.endswith('"')):
        return t[1:-1]
    # Numérico
    try:
        if "." in t:
            return float(t)
        return int(t)
    except Exception:
        return t


def _eval_condition(expr: str, ctx: Dict[str, Any]) -> bool:
    """
    Evalúa condiciones simples:
      VAR
      !VAR
      VAR == 'valor'
      VAR != "valor"
    """
    expr = expr.strip()

    # Negación simple !VAR
    if expr.startswith("!"):
        key = expr[1:].strip()
        return not _truthy(_get_ctx_value(ctx, key))

    # Igualdad / desigualdad
    m_eq = re.match(r"^([A-Za-z0-9_\.]+)\s*==\s*(.+)$", expr)
    if m_eq:
        key, lit = m_eq.group(1), m_eq.group(2)
        left = _get_ctx_value(ctx, key)
        right = _parse_literal(lit)
        return left == right

    m_neq = re.match(r"^([A-Za-z0-9_\.]+)\s*!=\s*(.+)$", expr)
    if m_neq:
        key, lit = m_neq.group(1), m_neq.group(2)
        left = _get_ctx_value(ctx, key)
        right = _parse_literal(lit)
        return left != right

    # Variable simple
    val = _get_ctx_value(ctx, expr)
    return _truthy(val)


# --------------------------------------------------------------------------------------
# Parser de condicionales con pila ([[IF]], [[ELIF]], [[ELSE]], [[/IF]])
# --------------------------------------------------------------------------------------
_TAG_RE = re.compile(r"\[\[(IF|ELIF|ELSE|/IF)(?:\s+(.*?))?\]\]", flags=re.IGNORECASE | re.DOTALL)

class _Node:
    __slots__ = ("children",)
    def __init__(self):
        self.children: List[Any] = []  # puede contener str o _IfNode

class _IfNode:
    __slots__ = ("clauses",)  # List[Tuple[Optional[str], _Node]]
    def __init__(self):
        self.clauses: List[Tuple[Optional[str], _Node]] = []  # (cond_expr or None for ELSE, node)


def _parse_to_ast(text: str) -> _Node:
    root = _Node()
    stack: List[_Node] = [root]
    cur = root
    last_pos = 0

    for m in _TAG_RE.finditer(text):
        # Texto previo al tag
        if m.start() > last_pos:
            cur.children.append(text[last_pos:m.start()])
        tag = m.group(1).upper()
        arg = (m.group(2) or "").strip()

        if tag == "IF":
            if_node = _IfNode()
            # primera cláusula (IF) arranca con su contenedor
            clause_node = _Node()
            if_node.clauses.append((arg, clause_node))
            # insertar el if_node en el árbol actual
            cur.children.append(if_node)
            # push: primero el if_node (lógico), pero las inserciones van al clause_node
            stack.append(cur)           # recordamos el padre
            stack.append(if_node)       # recordamos el if
            cur = clause_node           # trabajamos dentro de la 1ra cláusula
        elif tag == "ELIF":
            # cerramos cláusula previa y abrimos nueva dentro del mismo if
            if len(stack) < 2 or not isinstance(stack[-1], _IfNode):
                raise ValueError("Etiqueta [[ELIF]] sin [[IF]] correspondiente")
            if_node = stack[-1]
            clause_node = _Node()
            if_node.clauses.append((arg, clause_node))
            cur = clause_node
        elif tag == "ELSE":
            if len(stack) < 2 or not isinstance(stack[-1], _IfNode):
                raise ValueError("Etiqueta [[ELSE]] sin [[IF]] correspondiente")
            if_node = stack[-1]
            clause_node = _Node()
            if_node.clauses.append((None, clause_node))  # None = rama else
            cur = clause_node
        elif tag == "/IF":
            # cerramos el if actual y volvemos al padre
            if len(stack) < 2 or not isinstance(stack[-1], _IfNode):
                raise ValueError("Etiqueta [[/IF]] sin [[IF]] correspondiente")
            stack.pop()           # quitamos el if_node
            parent = stack.pop()  # recuperamos el padre (un _Node)
            cur = parent
        else:
            # No debería ocurrir
            pass

        last_pos = m.end()

    # resto del texto
    if last_pos < len(text):
        cur.children.append(text[last_pos:])

    # Si quedaron ifs sin cerrar, ignoramos lo que falte pero avisamos
    # (no interrumpimos para mantener robustez)
    return root


def _eval_ast(node: _Node, ctx: Dict[str, Any]) -> str:
    out_parts: List[str] = []
    for child in node.children:
        if isinstance(child, str):
            out_parts.append(child)
        elif isinstance(child, _IfNode):
            # evaluar cláusulas en orden
            chosen: Optional[_Node] = None
            for cond, clause_node in child.clauses:
                if cond is None:  # ELSE
                    chosen = clause_node
                    # no break; solo se usa si ninguna cond fue True
                else:
                    try:
                        if _eval_condition(cond, ctx):
                            chosen = clause_node
                            break
                    except Exception:
                        # condición mal formada -> se ignora como False
                        continue
            if chosen is not None:
                out_parts.append(_eval_ast(chosen, ctx))
            # si no hay rama elegida y no hay ELSE, no agrega nada
        else:
            # tipo desconocido -> nada
            continue
    return "".join(out_parts)


def render_conditionals(text: str, ctx: Dict[str, Any]) -> str:
    """
    Renderiza condicionales con anidamiento correcto y soporte de [[ELIF]]/[[ELSE]].
    """
    ast = _parse_to_ast(text)
    return _eval_ast(ast, ctx)


# --------------------------------------------------------------------------------------
# Sustitución de variables con filtros simples
# --------------------------------------------------------------------------------------
# {{ var }} o {{ var|upper }} |lower |title
_VAR_RE = re.compile(r"\{\{\s*([A-Za-z0-9_\.]+)\s*(?:\|\s*(upper|lower|title))?\s*\}\}")

def _apply_filter(val: Any, filt: Optional[str]) -> str:
    if val is None:
        return ""
    s = str(val)
    if not filt:
        return s
    if filt == "upper":
        return s.upper()
    if filt == "lower":
        return s.lower()
    if filt == "title":
        return s.title()
    return s


def render_variables(text: str, ctx: Dict[str, Any]) -> str:
    def repl(m: re.Match) -> str:
        key = m.group(1)
        filt = m.group(2)
        val = _get_ctx_value(ctx, key)
        return _apply_filter(val, filt)
    return _VAR_RE.sub(repl, text)


# --------------------------------------------------------------------------------------
# Render completo
# --------------------------------------------------------------------------------------
def _clean_text(out: str) -> str:
    # Eliminar espacios al final de línea
    out = re.sub(r"[ \t]+$", "", out, flags=re.MULTILINE)
    # Colapsar líneas en blanco (>2 seguidas -> 2)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out


# Patrón para detectar restos de placeholders/condicionales
_RE_VAR_LEFT = re.compile(r"\{\{\s*[A-Za-z0-9_\.]+(?:\s*\|\s*(?:upper|lower|title))?\s*\}\}")
_RE_TAG_LEFT = re.compile(r"\[\[(?:IF|ELIF|ELSE|/IF)[^\]]*\]\]", flags=re.IGNORECASE)

def render_text(raw_template: str, ctx: Dict[str, Any], estricto: bool = False) -> str:
    """
    1) Evalúa condicionales (con anidamiento real)
    2) Sustituye variables {{...}} con filtros simples
    3) Limpia y valida (modo estricto)
    """
    out = render_conditionals(raw_template, ctx)
    out = render_variables(out, ctx)
    out = _clean_text(out)

    if estricto:
        leftovers = []
        if _RE_VAR_LEFT.search(out):
            leftovers.append("placeholders {{...}}")
        if _RE_TAG_LEFT.search(out):
            leftovers.append("etiquetas [[IF/ELIF/ELSE/IF]]")
        if leftovers:
            raise ValueError("Quedaron marcadores sin resolver: " + ", ".join(leftovers))
    return out


def render_docx(plantilla: Path, ctx: Dict[str, Any], out_path: Path, estricto: bool = False) -> None:
    raw = doc_to_text(plantilla)
    final_text = render_text(raw, ctx, estricto=estricto)
    text_to_doc(final_text, out_path)


# --------------------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------------------
def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Render de descargo .docx a partir de plantilla y JSON (v2, con IF/ELIF/ELSE anidados)."
    )
    p.add_argument(
        "--plantilla",
        type=str,
        default=str(DIR_PLANTILLAS / "MODELO_DESCARGO_INTEGRAL_EXTENSO.docx"),
        help="Ruta a la plantilla .docx"
    )
    p.add_argument(
        "--caso",
        type=str,
        default=str(DIR_CASOS / "datos_ejemplo.json"),
        help="Ruta al JSON con datos del caso"
    )
    p.add_argument(
        "--salida",
        type=str,
        default="",  # si viene vacío, se arma con NRO_ACTA y fecha
        help="Ruta del .docx de salida (opcional)"
    )
    p.add_argument(
        "--estricto",
        action="store_true",
        help="Falla si quedan placeholders o etiquetas condicionales sin resolver."
    )
    return p.parse_args()


def _apply_defaults(ctx: Dict[str, Any]) -> Dict[str, Any]:
    """
    Defaults útiles para plantillas:
    - FECHA_PRESENTACION -> hoy (si no viene)
    """
    ctx = dict(ctx)  # copia superficial
    ctx.setdefault("FECHA_PRESENTACION", str(date.today()))
    return ctx


def main() -> None:
    ensure_dirs()
    args = parse_args()

    plantilla_path = Path(args.plantilla).resolve()
    caso_path = Path(args.caso).resolve()

    ctx = load_json(caso_path)
    ctx = _apply_defaults(ctx)

    # Nombre de salida por defecto
    if args.salida:
        out_path = Path(args.salida).resolve()
        out_path.parent.mkdir(parents=True, exist_ok=True)
    else:
        nro = ctx.get("NRO_ACTA", "SIN_ACTA")
        out_path = (DIR_SALIDAS / f"DESCARGO_{nro}_{date.today()}.docx").resolve()

    # Render
    try:
        render_docx(plantilla_path, ctx, out_path, estricto=args.estricto)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)

    print(f"OK -> {out_path}")


if __name__ == "__main__":
    main()
